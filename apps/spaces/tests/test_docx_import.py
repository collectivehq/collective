from __future__ import annotations

from io import BytesIO

import pytest
from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import Client
from django.urls import reverse
from docx import Document
from docx.shared import Pt

from apps.nodes.models import Node
from apps.spaces import services as space_services
from apps.spaces.docx_import import DocxImportError, import_space_from_docx, parse_space_docx
from apps.spaces.models import Space
from apps.users.tests.factories import UserFactory


def _add_structure_item(document: Document, label: str, *, level: int) -> None:
    if level == 0:
        document.add_paragraph(label, style="List Bullet")
        return
    if level == 1:
        document.add_paragraph(label, style="List Bullet 2")
        return
    if level == 2:
        document.add_paragraph(label, style="List Bullet 3")
        return
    if level == 3:
        paragraph = document.add_paragraph(label)
        paragraph.paragraph_format.left_indent = Pt(54)
        return
    raise ValueError(f"Unsupported structure level: {level}")


def _build_docx_bytes() -> bytes:
    document = Document()
    document.add_heading("Structure", level=1)
    _add_structure_item(document, "Q1", level=0)
    _add_structure_item(document, "1a", level=1)
    _add_structure_item(document, "1a-i", level=2)
    _add_structure_item(document, "1a-i-alpha", level=3)
    _add_structure_item(document, "1b", level=1)
    _add_structure_item(document, "Q2", level=0)

    document.add_heading("Content", level=1)
    document.add_heading("Q1", level=1)
    document.add_paragraph("Lorem ipsum")
    document.add_heading("1a", level=2)
    document.add_paragraph("Lalala")
    document.add_heading("1a-i", level=3)
    document.add_paragraph("Bababa")
    document.add_heading("1a-i-alpha", level=4)
    document.add_paragraph("Heading four content")
    document.add_heading("Q2", level=1)
    document.add_paragraph("Dolor sit amet")

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


@pytest.mark.django_db
class TestParseSpaceDocx:
    def test_parses_structure_and_content(self):
        parsed = parse_space_docx(docx_bytes=_build_docx_bytes())

        assert [node.label for node in parsed] == ["Q1", "Q2"]
        assert [child.label for child in parsed[0].children] == ["1a", "1b"]
        assert parsed[0].children[0].children[0].label == "1a-i"
        assert parsed[0].children[0].children[0].children[0].label == "1a-i-alpha"
        assert "Lorem ipsum" in parsed[0].content_html
        assert "Bababa" in parsed[0].children[0].children[0].content_html
        assert "Heading four content" in parsed[0].children[0].children[0].children[0].content_html

    def test_rejects_unknown_content_headings(self):
        document = Document()
        document.add_heading("Structure", level=1)
        document.add_paragraph("Q1", style="List Bullet")
        document.add_heading("Content", level=1)
        document.add_heading("Missing", level=1)
        document.add_paragraph("Oops")

        buffer = BytesIO()
        document.save(buffer)

        with pytest.raises(DocxImportError, match="unknown structure items"):
            parse_space_docx(docx_bytes=buffer.getvalue())


@pytest.mark.django_db
class TestImportSpaceFromDocx:
    def test_populates_discussions_and_posts(self):
        user = UserFactory()
        space = space_services.create_space(title="Imported", created_by=user)
        space_services.open_space(space=space)

        import_space_from_docx(space=space, author=user, docx_bytes=_build_docx_bytes())

        q1 = Node.objects.get(space=space, label="Q1", deleted_at__isnull=True, node_type=Node.NodeType.DISCUSSION)
        q2 = Node.objects.get(space=space, label="Q2", deleted_at__isnull=True, node_type=Node.NodeType.DISCUSSION)
        one_a = Node.objects.get(space=space, label="1a", deleted_at__isnull=True, node_type=Node.NodeType.DISCUSSION)
        one_a_i = Node.objects.get(
            space=space, label="1a-i", deleted_at__isnull=True, node_type=Node.NodeType.DISCUSSION
        )
        one_a_i_alpha = Node.objects.get(
            space=space,
            label="1a-i-alpha",
            deleted_at__isnull=True,
            node_type=Node.NodeType.DISCUSSION,
        )

        assert q1.get_parent().pk == space.root_discussion.pk
        assert q2.get_parent().pk == space.root_discussion.pk
        assert one_a.get_parent().pk == q1.pk
        assert one_a_i.get_parent().pk == one_a.pk
        assert one_a_i_alpha.get_parent().pk == one_a_i.pk
        assert Node.objects.filter(space=space, node_type=Node.NodeType.POST, content__contains="Lorem ipsum").exists()
        assert Node.objects.filter(space=space, node_type=Node.NodeType.POST, content__contains="Bababa").exists()
        assert Node.objects.filter(
            space=space,
            node_type=Node.NodeType.POST,
            content__contains="Heading four content",
        ).exists()


@pytest.mark.django_db
class TestSpaceCreateDocxImportView:
    def test_create_space_from_docx(self):
        user = UserFactory()
        client = Client()
        client.force_login(user)

        uploaded = SimpleUploadedFile(
            "structure.docx",
            _build_docx_bytes(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        response = client.post(
            reverse("spaces:create"),
            {
                "title": "Imported Space",
                "description": "Desc",
                "information": "",
                "source_docx": uploaded,
            },
        )

        assert response.status_code == 302
        space = Space.objects.get(title="Imported Space")
        assert Node.objects.filter(space=space, label="Q1", deleted_at__isnull=True).exists()
        assert Node.objects.filter(space=space, label="1a-i-alpha", deleted_at__isnull=True).exists()
        assert Node.objects.filter(
            space=space, node_type=Node.NodeType.POST, content__contains="Dolor sit amet"
        ).exists()
